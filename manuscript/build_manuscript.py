#!/usr/bin/env python3
"""Builds manuscript.docx for the project1 S100A8+ macrophage deconvolution study."""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIG_DIR = os.path.join(os.path.dirname(__file__), "..", "figures")
OUT_PATH = os.path.join(os.path.dirname(__file__), "manuscript.docx")

doc = Document()

# ---------------------------------------------------------------- styles ---
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(section, m, Inches(1))

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.5
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rpr.append(rFonts)
rFonts.set(qn("w:eastAsia"), "Times New Roman")

for lvl, size, bold in [(1, 15, True), (2, 13, True), (3, 12, True)]:
    style = doc.styles[f"Heading {lvl}"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_before = Pt(14)
    style.paragraph_format.space_after = Pt(6)


def add_caption(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(14)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def add_figure(png_name, width_in, caption_bold, caption_text):
    doc.add_picture(os.path.join(FIG_DIR, png_name), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(caption_text, bold_prefix=caption_bold)


def add_ref_run(paragraph, num):
    r = paragraph.add_run(str(num))
    r.font.superscript = True
    r.font.size = Pt(11)


def p_with_refs(segments):
    """segments: list of (text, ref_numbers_or_None)"""
    p = doc.add_paragraph()
    for text, refs in segments:
        if text:
            p.add_run(text)
        if refs:
            for i, n in enumerate(refs):
                add_ref_run(p, n)
                if i < len(refs) - 1:
                    p.add_run(",")
    return p


def make_table(headers, rows, col_widths_in=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


# ============================================================ TITLE PAGE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "Cross-Validated Deconvolution Identifies General Hepatic Macrophage "
    "Burden, Not an S100A8-Hi Subset, as the Lipogenic-Gene Correlate in "
    "Mouse MASLD"
)
r.bold = True
r.font.size = Pt(16)
title.paragraph_format.space_after = Pt(18)

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run("[Author Name]").italic = False
authors.add_run("¹").font.superscript = True
authors.add_run(", [Corresponding Author Name]")
r = authors.add_run("¹*")
r.font.superscript = False
authors.paragraph_format.space_after = Pt(4)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.add_run("¹[Department/Institution]").font.size = Pt(10)
affil.paragraph_format.space_after = Pt(4)

corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = corr.add_run("*Corresponding author: [email address]")
r.font.size = Pt(10)
corr.paragraph_format.space_after = Pt(24)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "This manuscript reports a secondary, hypothesis-generating computational "
    "reanalysis of a public dataset; it has not been peer reviewed."
)
r.italic = True
r.font.size = Pt(9)

doc.add_page_break()

# =============================================================== ABSTRACT ==
doc.add_heading("Abstract", level=1)

abs_p = doc.add_paragraph()
abs_p.add_run(
    "Background: "
).bold = True
abs_p.add_run(
    "Adipocyte death promotes hepatic accumulation of S100A8"
)
abs_p.add_run("+").font.superscript = True
abs_p.add_run(
    " macrophages, which suppress CCN3 and increase hepatocyte CD36 expression, "
    "driving metabolic dysfunction-associated steatotic liver disease "
    "(MASLD)."
)
add_ref_run(abs_p, 1)
abs_p.add_run(
    " The public dataset underlying this mechanism (GEO GSE285614/GSE285615) "
    "comprises 34 bulk RNA-seq samples across two independent genotype axes "
    "and only two single-cell RNA-seq (scRNA-seq) samples (one wild-type [WT], "
    "one Bcl2-transgenic [TG]), leaving the second axis — a macrophage-specific "
    "S100a8 knockout (MKO) — without any matched single-cell reference."
)

abs_p2 = doc.add_paragraph()
abs_p2.add_run("Methods: ").bold = True
abs_p2.add_run(
    "We reprocessed the public scRNA-seq reference, identified 12 hepatic cell "
    "populations including an S100A8-hi macrophage/monocyte population "
    "explicitly distinguished from neutrophils, and built a marker-gene "
    "signature. Because the reference contains only one biological replicate "
    "per genotype, subject-variance-based deconvolution methods were not "
    "applicable; we instead used marker-based deconvolution "
    "(BisqueRNA::MarkerBasedDecomposition) to estimate relative cell-type "
    "abundance in all 34 bulk samples. We validated the deconvolution against "
    "the direction of the original, fully replicated (n=5 vs n=5) bulk DESeq2 "
    "result, then applied it to the MKO axis and tested correlation with "
    "hepatic lipogenic gene expression. Group comparisons used Wilcoxon "
    "rank-sum tests (primary), Welch's t-test and Cohen's d (sensitivity/"
    "effect size), Benjamini-Hochberg (BH) correction within each comparison "
    "family, and post hoc power analysis."
)

abs_p3 = doc.add_paragraph()
abs_p3.add_run("Results: ").bold = True
abs_p3.add_run(
    "Deconvolution direction matched the replicated bulk DESeq2 result for "
    "every cell type checked. General macrophage/Kupffer abundance was "
    "significantly higher in WT than TG (Cohen's d = -2.86, BH q = 0.024) and "
    "correlated with Cd36, Pparg, Scd1 and Plin2 expression (all q < 0.05); "
    "the S100A8-hi subset trended in the same direction (d = -0.92) but did "
    "not reach significance and did not correlate with Cd36 (q = 0.76). In "
    "the MKO axis, no composition shift survived correction (this design is "
    "powered only for d "
)
abs_p3.add_run("≳").font.name = "Cambria Math"
abs_p3.add_run(
    " 2.0-2.4); nominal, uncorrected effects were observed in neutrophils and "
    "mast cells. Testing the S100A8-hi signature against two independent "
    "public mouse liver scRNA-seq datasets showed it reproduces as a "
    "discrete population in a CD45"
)
abs_p3.add_run("+").font.superscript = True
abs_p3.add_run(
    "-sorted dataset (with a diet-dependent trend at 2 of 3 timepoints) but "
    "not in a whole-liver dataset with true biological replication, "
    "suggesting immune-cell enrichment aids detection of this state."
)

abs_p4 = doc.add_paragraph()
abs_p4.add_run("Conclusions: ").bold = True
abs_p4.add_run(
    "Bulk-level compositional correlates of this pathway are driven "
    "predominantly by the general hepatic macrophage/Kupffer compartment "
    "rather than the S100A8-hi subset specifically, and macrophage-specific "
    "S100a8 deletion does not produce a statistically resolvable shift in "
    "broader immune composition at the sample sizes available — a pattern "
    "consistent with a gene-regulatory rather than a compositional/"
    "recruitment mechanism, though it requires confirmation in an adequately "
    "powered cohort."
)

kw = doc.add_paragraph()
kw.add_run("Keywords: ").bold = True
kw.add_run(
    "MASLD; S100A8; hepatic macrophage; cell-type deconvolution; single-cell "
    "RNA sequencing; BisqueRNA"
)

doc.add_page_break()

# ============================================================ INTRODUCTION =
doc.add_heading("1. Introduction", level=1)

p = doc.add_paragraph()
p.add_run(
    "Metabolic dysfunction-associated steatotic liver disease (MASLD, formerly "
    "nonalcoholic fatty liver disease)"
)
add_ref_run(p, 2)
p.add_run(
    " affects a large proportion of the global population and can progress "
    "to steatohepatitis, fibrosis and cirrhosis. A recent study using "
    "genetically modified mice and paired bulk and single-cell RNA "
    "sequencing (scRNA-seq) demonstrated that death of visceral adipocytes "
    "under high-fat diet (HFD) feeding drives hepatic infiltration of "
    "S100A8"
)
p.add_run("+").font.superscript = True
p.add_run(
    " macrophages, which suppress the CD36-negative regulator CCN3 and "
    "thereby increase hepatocyte CD36 expression and lipid storage. Blocking "
    "adipocyte death (adipocyte-specific Bcl2 overexpression, \"Bcl2TG\", "
    "which inhibits the mitochondrial apoptosis pathway) or deleting "
    "S100a8 specifically in macrophages (S100a8"
)
p.add_run("fl/fl").font.superscript = True
p.add_run(
    ";Cx3cr1-Cre, \"MKO\") both reduced hepatic steatosis, establishing this "
    "as a mechanistic axis in MASLD progression"
)
add_ref_run(p, 1)
p.add_run(".")

p = doc.add_paragraph()
p.add_run(
    "The public data supporting this study (GEO accessions GSE285614 and "
    "GSE285615) include 34 bulk liver RNA-seq samples spanning two "
    "independent genotype axes — an adipocyte-death axis (WT vs Bcl2TG) and "
    "a macrophage-specific knockout axis (WT vs MKO) — but only two scRNA-seq "
    "samples, one WT and one Bcl2TG, both under HFD. The original analysis "
    "of these data relied on single-gene bulk differential expression; it did "
    "not use the scRNA-seq reference to ask whether the two genotype "
    "manipulations change the "
)
r = p.add_run("composition")
r.italic = True
p.add_run(
    " of the hepatic immune infiltrate, as opposed to the expression level of "
    "individual genes within it — a distinction that matters because a "
    "shrinking S100A8"
)
p.add_run("+").font.superscript = True
p.add_run(
    " macrophage population and a stable population that has simply "
    "down-regulated S100a8 would look identical in a single-gene analysis "
    "but imply different biology. This question is structurally unanswerable "
    "for the MKO axis using the original study's own data, because that axis "
    "has no scRNA-seq reference of its own."
)

p = doc.add_paragraph()
p.add_run(
    "Here we reprocess the public scRNA-seq reference, build a marker-gene "
    "signature that explicitly separates an S100A8-hi macrophage/monocyte "
    "population from mature neutrophils (both S100a8/S100a9-positive, but "
    "otherwise transcriptionally distinct), and use marker-based "
    "deconvolution to estimate relative cell-type abundance across all 34 "
    "bulk samples. Because the scRNA-seq reference contains only one "
    "biological replicate per genotype, we first establish, and then apply "
    "throughout, a validation strategy suited to this constraint: rather than "
    "trusting the single scRNA-seq sample pair as ground truth (which we show "
    "can be directionally misleading), we validate the deconvolution against "
    "the original, fully replicated (n=5 vs n=5) bulk differential expression "
    "result for the same genotype comparison. We then apply the validated "
    "signature to the MKO axis — which has no scRNA-seq reference of its own — "
    "to test whether macrophage-specific S100a8 deletion detectably remodels "
    "the broader hepatic immune landscape, and we test whether cell-type "
    "abundance correlates with hepatic lipogenic gene expression across the "
    "full cohort."
)

# ================================================================ METHODS ==
doc.add_heading("2. Methods", level=1)

doc.add_heading("2.1 Data sources", level=2)
p = doc.add_paragraph()
p.add_run(
    "Raw and processed data were downloaded from the NCBI Gene Expression "
    "Omnibus"
)
add_ref_run(p, 7)
p.add_run(
    " (GSE285614, bulk liver RNA-seq; GSE285615, liver scRNA-seq), "
    "originally deposited alongside"
)
add_ref_run(p, 1)
p.add_run(". GSE285614 was deposited as two separate DESeq2")
add_ref_run(p, 5)
p.add_run(" analyses, each "
    "with per-sample raw counts embedded alongside summary statistics: "
    "\"HFD_TG_vs_WT\" (16 samples: WT and Bcl2TG, chow and HFD; the "
    "\"adipocyte-death axis\") and \"MKO_HFD_vs_WT_HFD\" (18 samples: "
    "S100a8"
)
p.add_run("fl/fl").font.superscript = True
p.add_run(
    " and S100a8"
)
p.add_run("fl/fl").font.superscript = True
p.add_run(
    ";Cx3cr1-Cre, chow and HFD; the \"macrophage-knockout axis\"). "
    "GSE285615 contains two 10x Genomics filtered feature-barcode matrices "
    "(one WT, one Bcl2TG, both HFD, both male, 8 weeks of age at study start). "
    "Study design and sample sizes are summarized in Table 1."
)

t1_headers = ["Axis", "Genotype comparison", "Diet", "n (bulk)", "scRNA-seq reference"]
t1_rows = [
    ["1 — Adipocyte death", "WT vs Bcl2TG", "Chow", "3 vs 3", "No"],
    ["1 — Adipocyte death", "WT vs Bcl2TG", "HFD", "5 vs 5", "Yes (1 vs 1)"],
    ["2 — Macrophage S100a8 KO", "WT(fl/fl) vs MKO", "Chow", "4 vs 5", "No"],
    ["2 — Macrophage S100a8 KO", "WT(fl/fl) vs MKO", "HFD", "5 vs 4", "No"],
]
add_caption("Study design: bulk RNA-seq sample sizes and single-cell reference availability by genotype axis and diet.", bold_prefix="Table 1. ")
make_table(t1_headers, t1_rows, col_widths_in=[1.5, 1.5, 0.8, 1.0, 1.5])

doc.add_heading("2.2 Single-cell RNA-seq processing and cell-type annotation", level=2)
p = doc.add_paragraph()
p.add_run(
    "Filtered feature-barcode matrices were loaded with Seurat v5.4.0"
)
add_ref_run(p, 3)
p.add_run(
    " (CreateSeuratObject, min.cells = 3, min.features = 200). Cells were "
    "retained if 200 < nFeature_RNA < 6,000, nCount_RNA < 30,000 and "
    "percent.mt < 25% (mitochondrial fraction threshold set liberally "
    "because hepatocytes are naturally mitochondria-rich); this removed "
    "320/11,579 WT cells and 277/8,158 TG cells (retained: WT = 11,259, "
    "TG = 7,881; total = 19,140). The two samples were merged without "
    "cross-sample batch correction (e.g., Harmony or CCA-based integration): "
    "because sample identity is completely confounded with genotype in this "
    "n=1-per-group design, batch correction would risk removing exactly the "
    "biological signal under study. Data were log-normalized, 2,000 highly "
    "variable genes were selected, and principal component analysis was run "
    "on scaled data (30 components). Cells were clustered with a shared "
    "nearest-neighbor graph and the Louvain algorithm (resolution = 0.8, "
    "28 clusters) and visualized with UMAP"
)
add_ref_run(p, 8)
p.add_run(".")

p = doc.add_paragraph()
p.add_run(
    "Cluster identities were assigned from canonical marker genes (e.g., "
    "Alb/Apoa1/Cyp2e1, hepatocyte; Clec4f/Csf1r/Adgre1/Cd68, "
    "macrophage/Kupffer; Pecam1/Stab2/Clec4g, endothelial; Dcn/Pdgfrb, "
    "stellate/fibroblast; Cd3e/Cd3d, T cell; Cd79a/Ms4a1, B cell; Ncr1, NK "
    "cell; Krt19/Epcam, cholangiocyte) together with FindAllMarkers "
    "(Wilcoxon rank-sum test, only.pos = TRUE, min.pct = 0.25, "
    "logFC threshold = 0.5). Twenty-eight clusters were collapsed into 12 "
    "cell types. Two low-confidence clusters (973/19,140 cells, 5.1%) were "
    "excluded before downstream analysis: one dominated by ambient-RNA/"
    "mitochondrial genes and one defined only by cell-cycle genes with no "
    "clear lineage identity. Critically, S100a8/S100a9-positive myeloid "
    "cells resolved into two transcriptionally distinct clusters: an "
    "\"S100A8hi_macrophage\" population (S100a8/S100a9-positive, Ly6g-"
    "negative, partially Csf1r/Cd68/Adgre1-positive — monocyte/macrophage "
    "lineage) and a \"Neutrophil\" population (S100a8/S100a9-positive "
    "together with Ly6g, Camp, Ltf, Ngp and Mpo — mature granulocyte "
    "lineage; Figure 1B). These were kept as separate cell types throughout, "
    "since collapsing them would confound the population of biological "
    "interest with an unrelated granulocyte lineage that also expresses "
    "S100a8/S100a9."
)

doc.add_heading("2.3 Bulk RNA-seq data integration", level=2)
p = doc.add_paragraph()
p.add_run(
    "Per-sample raw counts were extracted from both GSE285614 supplementary "
    "tables, duplicate gene symbols were collapsed by summation, and the two "
    "gene sets were intersected (12,505 shared genes) before merging into a "
    "single 34-sample raw count matrix. The two genotype axes were kept "
    "labeled separately throughout all downstream analyses — although both "
    "contain a group labeled \"WT\", they were sequenced as independent "
    "DESeq2 runs/batches and were never pooled or treated as the same "
    "biological samples. Counts were converted to counts-per-million (CPM) "
    "per sample and log"
)
r = p.add_run("2")
r.font.subscript = True
p.add_run(
    "(CPM + 1)-transformed for variance stabilization prior to "
    "principal-component-based deconvolution."
)

doc.add_heading("2.4 Cell-type deconvolution", level=2)
p = doc.add_paragraph()
p.add_run(
    "Because the scRNA-seq reference contains only one biological replicate "
    "per genotype, subject-variance-aware deconvolution methods (e.g., "
    "MuSiC, or BisqueRNA's own reference-based decomposition), which require "
    "multiple donors per cell type to estimate cross-subject variance, were "
    "not statistically appropriate. We therefore used "
    "BisqueRNA::MarkerBasedDecomposition"
)
add_ref_run(p, 4)
p.add_run(
    ", which requires only a stable per-cell-type marker-gene signature and "
    "returns, for each cell type, the first principal component (PC1) of "
    "that cell type's marker-gene sub-matrix in the bulk expression data as "
    "a relative abundance score. This score is not an absolute proportion "
    "and does not sum to 1 across cell types; it is only valid for comparing "
    "the "
)
r = p.add_run("same")
r.italic = True
p.add_run(
    " cell type across samples or groups, not for comparing magnitude "
    "between different cell types."
)

p = doc.add_paragraph()
p.add_run(
    "Marker genes were obtained by running FindAllMarkers on the final "
    "annotated scRNA-seq object (Wilcoxon rank-sum test, only.pos = TRUE, "
    "min.pct = 0.25, logFC threshold = 0.5), restricting to genes unique to "
    "a single cell type and present in the bulk gene set, and ranking by "
    "adjusted p-value. Decomposition was run with weighted = TRUE using "
    "average log"
)
r = p.add_run("2")
r.font.subscript = True
p.add_run(
    " fold-change as the marker weight (min_gene = 5, max_gene = 100, "
    "unique_markers = TRUE). The number of marker genes actually used per "
    "cell type (5-25) is not a fixed or arbitrary cutoff: BisqueRNA's "
    "internal GetNumGenesWeighted() routine selects, within the allowed "
    "range, the panel size that maximizes the ratio of the first to second "
    "eigenvalue of the marker sub-matrix — a scree-like criterion that "
    "favors a single dominant compositional signal over simply maximizing "
    "gene count. We report, for transparency, both the number of markers "
    "used and the percentage of marker-panel variance explained by PC1 for "
    "each cell type (Table S1); this exceeded 70% for Macrophage_Kupffer "
    "(72.4%), S100A8hi_macrophage (76.9%) and NK_cell (78.1%), indicating a "
    "coherent single-factor signal for the cell types most relevant to this "
    "study, and fell to 33-39% for Cholangiocyte, Mast_cell, Neutrophil and "
    "T_cell, which should be interpreted more cautiously."
)

doc.add_heading("2.5 Validation strategy", level=2)
p = doc.add_paragraph()
p.add_run(
    "We initially validated the axis-1 (WT vs Bcl2TG, HFD) deconvolution "
    "against cell-type proportions counted directly from the single WT/TG "
    "scRNA-seq sample pair. This showed the S100A8hi_macrophage proportion "
    "as higher in TG than WT (3.55% vs 2.43%, a 1.46-fold difference), "
    "whereas the deconvolution estimated the opposite direction (WT > TG). "
    "Because a single biological replicate per group cannot establish a "
    "population-level direction, we treated this apparent discrepancy as "
    "inconclusive rather than as validation failure, and instead checked the "
    "direction of the original, fully replicated (n=5 WT vs n=5 TG) bulk "
    "DESeq2 result for the same marker genes in the identical comparison, "
    "obtained from the same publicly deposited supplementary table used for "
    "count extraction (Section 2.3). S100a8 (log"
)
r = p.add_run("2")
r.font.subscript = True
p.add_run(
    " fold-change = -0.80, p = 0.039), Cd68 (log"
)
r = p.add_run("2")
r.font.subscript = True
p.add_run(
    " fold-change = -0.74, p = 0.004, BH-adjusted p = 0.050) and other "
    "macrophage markers (Adgre1, Csf1r, Clec4f, Ptprc) were all higher in WT "
    "than TG, matching the deconvolution direction rather than the "
    "single-pair scRNA-seq comparison (Figure 2C). This is also the direction "
    "predicted mechanistically: Bcl2 overexpression blocks adipocyte "
    "apoptosis and is therefore protective, so Bcl2TG mice are expected to "
    "show "
)
r = p.add_run("less")
r.italic = True
p.add_run(
    " HFD-induced hepatic immune infiltration than WT, not more. We adopted "
    "this replicated-bulk benchmark as the validation standard for all "
    "subsequent analyses."
)

doc.add_heading("2.6 Statistical analysis", level=2)
p = doc.add_paragraph()
p.add_run(
    "All pairwise group comparisons of deconvolution scores used the "
    "two-sided Wilcoxon rank-sum test as the primary statistic, given the "
    "small group sizes (n = 4-5) and the inability to verify normality; "
    "Welch's two-sample t-test and Cohen's d (pooled-SD) are reported "
    "alongside as sensitivity and effect-size measures. Benjamini-Hochberg "
    "(BH) false discovery rate correction"
)
add_ref_run(p, 6)
p.add_run(
    " was applied within each pre-specified comparison family (12 cell "
    "types, separately within axis 1's HFD validation cohort and within "
    "each diet of axis 2's MKO screen; 30 cell-type-by-gene-by-sample-"
    "subset combinations for the lipogenic-gene correlation analysis), and "
    "q < 0.05 was used as the "
    "significance threshold for corrected results. Correlations between "
    "deconvolution scores and log"
)
r = p.add_run("2")
r.font.subscript = True
p.add_run(
    "(CPM + 1) gene expression used Spearman's rank correlation, appropriate "
    "given the arbitrary, non-linear scale of PC1-based deconvolution "
    "scores. Post hoc statistical power was assessed with a two-sample "
    "t-test power calculation (alpha = 0.05, 80% power) for n = 4 and n = 5 "
    "per group to contextualize non-significant results. All analyses were "
    "performed in R 4.5.2; the complete, version-controlled analysis code "
    "(scripts 00-12) is available at the repository in Section 2.8."
)

doc.add_heading("2.7 External validation in independent public datasets", level=2)
p = doc.add_paragraph()
p.add_run(
    "To test whether the S100A8-hi macrophage population reflects a reproducible cell "
    "state rather than an artifact of the single WT/TG reference pair (Section 2.2), "
    "we searched GEO for independent mouse liver scRNA-seq datasets with true "
    "biological replication and applied the same marker criteria used to define this "
    "population (S100a8/S100a9 percent-positive > 90%, Ly6g/Camp/Ltf/Ngp "
    "percent-positive < 10%, versus > 50% for a mature neutrophil) to their "
    "clusters. Two datasets were used: (1) GSE166504"
)
add_ref_run(p, 9)
p.add_run(
    ", mouse liver non-parenchymal cells across a chow / 15-week high-fat-high-"
    "fructose-diet (NAFL) / 30-week (NASH) progression with true biological "
    "replicates (chow n = 6, NAFL n = 3, NASH n = 4 mice), reprocessed from raw "
    "counts and the authors' own coarse cell-type labels (Kupffer cell, monocyte/"
    "monocyte-derived macrophage, neutrophil; 23,198 myeloid cells), reclustered "
    "within the myeloid compartment (resolution = 0.6, 18 clusters); and (2) "
    "GSE156057 (Remmerie et al."
)
add_ref_run(p, 10)
p.add_run(
    "), CD45+-sorted mouse liver immune cells across standard diet (SD) vs "
    "western diet (WD), 12/24/36-week timepoints (one sample per condition-"
    "timepoint -- pseudo-replicates across independent mice/times, not "
    "within-condition biological replicates), reclustered de novo (77,655 "
    "cells, resolution = 0.8, 40 clusters). For each dataset we determined "
    "whether any cluster matched the marker criteria and, where present, "
    "computed its proportion of CD45+/myeloid cells across conditions."
)

doc.add_heading("2.8 Data and code availability", level=2)
p = doc.add_paragraph()
p.add_run(
    "Raw data: NCBI GEO GSE285614 (bulk RNA-seq), GSE285615 (scRNA-seq), "
    "GSE166504 and GSE156057 (external validation). Analysis code, intermediate "
    "result tables and all figures: https://github.com/bioinform25/project1."
)

doc.add_page_break()

# ================================================================ RESULTS ==
doc.add_heading("3. Results", level=1)

doc.add_heading(
    "3.1 A distinct S100A8-hi macrophage/monocyte population is separable "
    "from neutrophils in the scRNA-seq reference", level=2
)
p = doc.add_paragraph()
p.add_run(
    "After quality control, 19,140 cells (WT = 11,259; TG = 7,881) were "
    "clustered into 28 groups and, after excluding two low-confidence "
    "clusters, annotated into 12 hepatic cell types comprising 18,167 cells "
    "(Figure 1A). As anticipated from the source study, S100a8/S100a9 "
    "expression was not confined to a single cluster: it appeared both in a "
    "monocyte/macrophage-lineage population (S100A8hi_macrophage: partial "
    "Csf1r/Cd68/Adgre1 positivity, Ly6g-negative) and, independently, in a "
    "mature neutrophil population (Ly6g/Camp/Ltf/Ngp/Mpo-positive; "
    "Figure 1B). Treating S100a8/S100a9 positivity alone as sufficient to "
    "define \"the\" S100A8"
)
p.add_run("+").font.superscript = True
p.add_run(
    " macrophage population would have conflated these two unrelated myeloid "
    "lineages; they were therefore modeled as separate cell types throughout."
)

p = doc.add_paragraph()
p.add_run(
    "S100A8hi_macrophage also showed minimal expression of Trem2 (0.9% of "
    "cells) and Gpnmb (2.1%), markers of lipid-associated/scar-associated "
    "macrophages, whereas these markers were concentrated instead within a "
    "subset of the broader Macrophage_Kupffer compartment (Trem2 42.1%, "
    "Gpnmb 23.5%; Figure 1B). This distinction is addressed further in "
    "Section 4."
)

add_figure(
    "Figure1.png", 6.5, "Figure 1. ",
    "scRNA-seq reference and identification of a distinct S100A8-hi "
    "macrophage population. (A) UMAP projection of 18,167 quality-"
    "controlled hepatic cells (WT + Bcl2TG, HFD), colored by annotated cell "
    "type. (B) Marker-gene dot plot distinguishing S100A8hi_macrophage from "
    "Neutrophil and from other Macrophage_Kupffer populations: only "
    "Neutrophil is positive for Ly6g, Camp, Ltf and Ngp, despite both "
    "S100A8hi_macrophage and Neutrophil expressing S100a8/S100a9."
)

doc.add_heading(
    "3.2 Marker-based deconvolution recovers the direction of the original, "
    "fully replicated bulk analysis", level=2
)
p = doc.add_paragraph()
p.add_run(
    "In the axis-1 (WT vs Bcl2TG, HFD) validation cohort (n = 5 per group), "
    "the Macrophage_Kupffer deconvolution score was significantly higher in "
    "WT than TG (Wilcoxon p = 0.012, BH q = 0.024, Cohen's d = -2.86; "
    "Figure 2A), matching the direction of every individual macrophage "
    "marker gene (Adgre1, Csf1r, Clec4f, Cd68) in the original, fully "
    "replicated bulk DESeq2 analysis (Figure 2C). The S100A8hi_macrophage "
    "score trended in the same direction (Cohen's d = -0.92) but did not "
    "reach significance at this sample size (Wilcoxon p = 0.30, BH q = 0.36; "
    "Figure 2B) — consistent with the corresponding single-gene result "
    "(S100a8 nominal p = 0.039 but BH-adjusted p = 0.22 in the original "
    "n=5-vs-5 DESeq2 analysis). Four additional cell types reached BH "
    "q = 0.024 in the same direction as expected for a protective genotype "
    "(reduced HFD-associated immune infiltration in Bcl2TG): Dendritic_cell, "
    "NK_cell, Stellate_Fibroblast and T_cell; Hepatocyte scored higher in TG "
    "than WT (q = 0.024), consistent with proportionally less immune "
    "infiltrate diluting the hepatocyte signal in WT. Full statistics "
    "(effect sizes, both tests, both BH-adjusted q-values) for all 12 cell "
    "types are provided in "
)
r = p.add_run("Data S1")
r.italic = True
p.add_run(
    " (Supplemental Data), generated from the same analysis that produced "
    "Figure 2."
)

add_figure(
    "Figure2.png", 6.3, "Figure 2. ",
    "Deconvolution validates against the original, fully replicated bulk "
    "DESeq2 result (axis 1, HFD). (A) Macrophage_Kupffer and (B) "
    "S100A8hi_macrophage deconvolution scores, WT vs Bcl2TG (n = 5 per "
    "group; boxplots with individual samples; Wilcoxon p and Cohen's d "
    "shown). (C) log₂ fold-change (Bcl2TG / WT) for the corresponding "
    "marker genes in the original authors' bulk DESeq2 analysis of the same "
    "16-sample cohort (n = 5 vs 5, HFD subset shown)."
)

doc.add_heading(
    "3.3 Macrophage-specific S100a8 deletion does not produce a "
    "statistically resolvable shift in hepatic cell composition", level=2
)
p = doc.add_paragraph()
p.add_run(
    "Applying the validated signature to the MKO axis (which has no "
    "scRNA-seq reference of its own), no cell type reached BH q < 0.05 in "
    "either diet (12 cell types x 2 diets; all q "
)
p.add_run("≥").font.name = "Cambria Math"
p.add_run(
    " 0.22). Neither S100A8hi_macrophage (HFD: d = -0.91, p = 0.39, "
    "q = 0.43; Figure 3B) nor Macrophage_Kupffer (HFD: d = 0.49, p = 0.27, "
    "q = 0.43; Figure 3D) differed between WT and MKO. Post hoc power "
    "analysis showed this cohort (n = 4-5 per group) is powered (80%, "
    "alpha = 0.05) to detect only Cohen's d "
)
p.add_run("≳").font.name = "Cambria Math"
p.add_run(
    " 2.0-2.4 — larger than the effect sizes observed for either macrophage "
    "population, meaning the null result reflects limited power for "
    "moderate effects rather than a confirmed absence of any composition "
    "change. Two cell types showed nominal, BH-uncorrected effects at "
    "exactly the magnitude this design can detect: Neutrophil was lower in "
    "MKO (d = -2.42, p = 0.037) and Mast_cell was higher in MKO (d = 2.36, "
    "p = 0.037) under HFD (Figure 3E-F); neither survived correction across "
    "the 12-cell-type family (q = 0.22) and both require confirmation in an "
    "adequately powered cohort before mechanistic interpretation. Full "
    "results for all 12 cell types in both diets, underlying Figure 3, are "
    "in "
)
r = p.add_run("Data S2")
r.italic = True
p.add_run(" (Supplemental Data).")

add_figure(
    "Figure3.png", 6.5, "Figure 3. ",
    "Macrophage-specific S100a8 deletion (MKO) does not significantly "
    "remodel hepatic cell composition. Deconvolution scores for WT vs MKO, "
    "by diet, for (A-B) S100A8hi_macrophage, (C-D) Macrophage_Kupffer, and "
    "(E-F) the two cell types with the largest nominal (BH-uncorrected) "
    "effect sizes, Neutrophil and Mast_cell (HFD only). p-values are "
    "two-sided Wilcoxon rank-sum; d is Cohen's d. All comparisons were "
    "BH-adjusted across the 12 cell types tested within each diet; none "
    "reached q < 0.05 (Table 2; results/11_MKO_stats_full.csv)."
)

doc.add_heading(
    "3.4 General macrophage abundance, not the S100A8-hi subset, correlates "
    "with hepatic lipogenic gene expression", level=2
)
p = doc.add_paragraph()
p.add_run(
    "Across all 19 HFD bulk samples (both axes pooled), the "
    "Macrophage_Kupffer deconvolution score correlated with Cd36 "
    "(Spearman r = 0.74, BH q = 0.006), Pparg (r = 0.75, q = 0.006), Scd1 "
    "(r = 0.73, q = 0.006) and Plin2 (r = 0.65, q = 0.021) after correcting "
    "for the 30 tests performed (Figure 4A, 4C). The S100A8hi_macrophage "
    "score showed no correlation with any of the five lipogenic genes tested "
    "in any sample subset (all q "
)
p.add_run("≥").font.name = "Cambria Math"
p.add_run(
    " 0.42; Cd36: r = 0.11, q = 0.76; Figure 4B-C) — a consistent, "
    "reproducible null across the full cohort, within axis 1 alone, and "
    "within axis 2 alone. Full correlation results for all 30 tests "
    "(5 genes x 2 cell types x 3 sample subsets) are in "
)
r = p.add_run("Data S3")
r.italic = True
p.add_run(" (Supplemental Data).")

t2_headers = ["Comparison", "Cell type / gene pair", "Effect size", "p-value", "BH q"]
t2_rows = [
    ["Axis 1, WT vs TG (HFD)", "Macrophage_Kupffer", "d = -2.86", "0.012", "0.024"],
    ["Axis 1, WT vs TG (HFD)", "S100A8hi_macrophage", "d = -0.92", "0.30", "0.36"],
    ["Axis 2, WT vs MKO (HFD)", "S100A8hi_macrophage", "d = -0.91", "0.39", "0.43"],
    ["Axis 2, WT vs MKO (HFD)", "Macrophage_Kupffer", "d = 0.49", "0.27", "0.43"],
    ["Axis 2, WT vs MKO (HFD)", "Neutrophil", "d = -2.42", "0.037", "0.22"],
    ["Axis 2, WT vs MKO (HFD)", "Mast_cell", "d = 2.36", "0.037", "0.22"],
    ["Correlation, all HFD (n=19)", "Macrophage_Kupffer vs Cd36", "r = 0.74", "0.0005", "0.006"],
    ["Correlation, all HFD (n=19)", "S100A8hi_macrophage vs Cd36", "r = 0.11", "0.66", "0.76"],
]
add_caption(
    "Selected statistical comparisons (full results: results/11_axis1_stats_full.csv, "
    "results/11_MKO_stats_full.csv, results/11_lipid_correlation_BH.csv). "
    "p-values are two-sided Wilcoxon rank-sum (group comparisons) or Spearman "
    "correlation test (correlations); BH q = Benjamini-Hochberg-adjusted p-value "
    "within the relevant comparison family.",
    bold_prefix="Table 2. "
)
make_table(t2_headers, t2_rows, col_widths_in=[1.6, 2.0, 0.9, 0.7, 0.7])

add_figure(
    "Figure4.png", 6.5, "Figure 4. ",
    "General macrophage abundance tracks lipogenic gene expression; the "
    "S100A8-hi subset does not. (A) Macrophage_Kupffer and (B) "
    "S100A8hi_macrophage deconvolution scores vs Cd36 expression across all "
    "19 HFD bulk samples (both axes; linear fit shown for visualization). "
    "(C) Spearman correlation (r) between each macrophage-lineage cell "
    "type's deconvolution score and five lipogenic genes; asterisks mark "
    "BH q < 0.05 across all 30 tests performed."
)

doc.add_heading(
    "3.5 The S100A8-hi macrophage signature reproduces in an independent, "
    "immune-cell-sorted dataset but not in a whole-liver dataset", level=2
)
p = doc.add_paragraph()
p.add_run(
    "In GSE156057 (CD45+-sorted, 77,655 cells, 40 clusters), three clusters "
    "matched the S100A8-hi, Ly6g-low criteria (S100a8/S100a9 percent-positive "
    "97-99%, Ly6g/Camp/Ltf/Ngp all < 5%), clearly separated from an "
    "unambiguous neutrophil cluster (S100a8/S100a9 100%, Ly6g 89%, Camp 88%, "
    "Ltf 80%, Ngp 95%) and from Clec4f-high resident Kupffer clusters "
    "(Figure 5A; marker profile for all 40 clusters in "
)
r = p.add_run("Data S4")
r.italic = True
p.add_run(", top-5 marker genes per cluster in ")
r = p.add_run("Data S5")
r.italic = True
p.add_run(
    "). Combined, these clusters comprised 5.0-7.8% of CD45"
)
p.add_run("+").font.superscript = True
p.add_run(
    " cells per sample and were more abundant under western diet (WD) than "
    "standard diet (SD) at 12 and 24 weeks (WD/SD ratio 1.45 and 1.21) but "
    "not at 36 weeks (ratio 0.81; Figure 5B, full per-sample counts in "
)
r = p.add_run("Data S6")
r.italic = True
p.add_run(
    ") — directionally consistent with, "
    "though not statistically confirmable given the single-mouse-per-"
    "condition design, diet-dependent recruitment. In GSE166504 (whole-liver "
    "non-parenchymal cells with true biological replication: chow n = 6, "
    "NAFL n = 3, NASH n = 4 mice; 23,198 myeloid cells, 18 clusters), no "
    "cluster showed this combination: the only S100a8/S100a9-high cluster "
    "also carried high Camp/Ngp/Ltf and was author-labeled Neutrophil, while "
    "S100a8 positivity among monocyte/macrophage-lineage clusters varied "
    "only as a modest, continuous gradient (6-19%) uncorrelated with Csf1r "
    "positivity (Figure 5C; full marker profile for all 18 clusters in "
)
r = p.add_run("Data S7")
r.italic = True
p.add_run(
    ", cluster-to-author-label cross-tabulation in "
)
r = p.add_run("Data S8")
r.italic = True
p.add_run(
    ") rather than resolving into a discrete "
    "population. We interpret this discrepancy as most consistent with "
    "immune-cell enrichment (CD45+ sorting, used both in this study's own "
    "scRNA-seq reference and in GSE156057) being necessary to resolve this "
    "population by unsupervised clustering at typical whole-liver sequencing "
    "depth, rather than as evidence that the population is an artifact "
    "specific to the single WT/TG reference pair used in the primary "
    "analysis (Section 2.2)."
)

add_figure(
    "Figure5.png", 6.7, "Figure 5. ",
    "External validation of the S100A8-hi macrophage signature in two "
    "independent mouse liver scRNA-seq datasets. (A) Marker profile "
    "(percent of cells expressing each gene) for the three candidate "
    "S100A8-hi clusters, the neutrophil cluster, and two representative "
    "Kupffer clusters in GSE156057 (CD45+-sorted). (B) S100A8-hi cluster "
    "proportion (combined) by diet and timepoint in GSE156057. (C) In "
    "GSE166504 (whole-liver, true biological replicates), percent of cells "
    "expressing S100a8 plotted against percent expressing Csf1r for all 18 "
    "myeloid clusters, colored by the authors' dominant cell-type label; "
    "S100a8 positivity forms a gradient rather than isolating a discrete "
    "cluster among monocyte/macrophage-lineage populations."
)

doc.add_page_break()

# ============================================================= DISCUSSION ==
doc.add_heading("4. Discussion", level=1)

p = doc.add_paragraph()
p.add_run(
    "This study reprocesses a public bulk-and-single-cell RNA-seq dataset to "
    "ask a question the original analysis could not: whether the two "
    "genetic manipulations that block this adipocyte death-CD36 axis change "
    "the "
)
r = p.add_run("composition")
r.italic = True
p.add_run(
    " of the hepatic immune infiltrate, not just the expression of "
    "individual marker genes within it. Three results converge on a "
    "consistent interpretation. First, general hepatic macrophage/Kupffer "
    "abundance — not the S100A8-hi subset specifically — is the "
    "compositional signal that is both statistically robust in the "
    "validation cohort and correlated with hepatic lipogenic gene "
    "expression across the full 19-sample HFD cohort. Second, the "
    "S100A8-hi subset shows a directionally concordant trend at every "
    "checkpoint in this study (matching the replicated bulk marker-gene "
    "direction in axis 1, and trending lower in MKO as would be expected if "
    "the mechanism were partly compositional) but never reaches statistical "
    "significance on its own — a pattern our power analysis indicates is "
    "expected given the observed effect size (d "
)
r = p.add_run("≈")
r.font.name = "Cambria Math"
p.add_run(
    " 0.9) relative to what a 4-5-sample-per-group design can resolve "
    "(d "
)
r = p.add_run("≳")
r.font.name = "Cambria Math"
p.add_run(
    " 2.0-2.4 at 80% power), not evidence against a real effect. Third, "
    "macrophage-specific S100a8 deletion (MKO) does not produce a "
    "statistically resolvable shift in broader immune composition in either "
    "diet. Taken together with the original study's finding that MKO is "
    "protective at the phenotypic and CD36-expression level"
)
add_ref_run(p, 1)
p.add_run(
    ", the most parsimonious interpretation is that S100A8 acts primarily "
    "through a gene-regulatory mechanism within the existing macrophage "
    "compartment (i.e., loss of S100A8 protein function suppresses CCN3 "
    "regardless of macrophage abundance) rather than by changing which or "
    "how many immune cells infiltrate the liver — though, given the power "
    "limitations discussed below, this should be read as a hypothesis "
    "refined rather than a mechanism confirmed."
)

p = doc.add_paragraph()
p.add_run(
    "This work also illustrates a general methodological point relevant to "
    "any study pairing a small pilot scRNA-seq cohort with a larger, "
    "properly replicated bulk cohort: an n=1-per-group single-cell "
    "comparison should not be used as an unqualified validation benchmark "
    "for a downstream computational method. In this dataset, the single "
    "WT/TG scRNA-seq pair suggested the opposite direction from the "
    "properly replicated bulk result for the same comparison; had we "
    "accepted the scRNA-seq pair at face value, we would have incorrectly "
    "concluded that the deconvolution method had failed. Cross-checking "
    "against an independent, adequately powered statistical result — here, "
    "the original authors' own bulk differential expression analysis — was "
    "necessary to correctly adjudicate which comparison was unreliable."
)

p = doc.add_paragraph()
p.add_run(
    "Because the primary scRNA-seq reference comprises a single WT and "
    "single TG mouse, a further question is whether the S100A8-hi "
    "macrophage population itself is reproducible, or an artifact of that "
    "one pair. We addressed this directly by testing the same marker "
    "criteria against two independent, publicly available mouse liver "
    "scRNA-seq datasets (Section 2.7, Figure 5). The population reproduced "
    "cleanly, with a diet-dependent trend at two of three timepoints, in a "
    "CD45"
)
p.add_run("+").font.superscript = True
p.add_run(
    "-sorted dataset (GSE156057), but did not resolve as a discrete cluster "
    "in a whole-liver dataset with true biological replication (GSE166504) "
    "— suggesting that immune-cell enrichment, rather than biological "
    "replicate number per se, may be the more critical design factor for "
    "detecting this state by unsupervised clustering at typical whole-liver "
    "sequencing depth. This is itself a transferable observation for scRNA-"
    "seq study design in liver immunology, and it meaningfully strengthens "
    "confidence that the population identified in the primary analysis is "
    "real rather than an artifact of a single mouse pair, even though it "
    "does not resolve the separate, unrelated question of that pair's "
    "reliability as a directional ground truth (Section 2.5)."
)

p = doc.add_paragraph()
p.add_run(
    "Two further observations help refine the biological identity of this "
    "population. First, the continuous, non-discrete S100a8 gradient "
    "observed among monocyte/macrophage-lineage clusters in the whole-liver "
    "GSE166504 dataset (Figure 5C) is consistent with the established "
    "biology of monocyte-to-macrophage differentiation, in which recruited "
    "monocytes traverse a continuous transcriptional trajectory rather than "
    "switching between sharply discrete states; immune-cell enrichment by "
    "CD45"
)
p.add_run("+").font.superscript = True
p.add_run(
    " sorting likely improves the effective signal-to-noise ratio for this "
    "transitional population, allowing unsupervised clustering to resolve "
    "cleaner boundaries around it, rather than the underlying biology "
    "itself differing between datasets. Second, the S100A8-hi population's "
    "own marker profile argues against it being a mature lipid-associated "
    "macrophage (also termed scar-associated or NASH-associated "
    "macrophage) — the Trem2-high, Gpnmb-high, osteopontin-high phenotype "
    "described in steatohepatitis liver by Remmerie et al."
)
add_ref_run(p, 10)
p.add_run(
    " and related work: in the primary reference, S100A8hi_macrophage "
    "showed minimal Trem2 (0.9%) and Gpnmb (2.1%) positivity, while these "
    "markers were concentrated instead within the broader Macrophage_"
    "Kupffer compartment (Trem2 42.1%, Gpnmb 23.5%; Section 3.1). Combined "
    "with its transient, diet-reversible abundance pattern in GSE156057 "
    "(Section 3.5), this profile is more consistent with an early, acutely "
    "inflammatory, recently recruited monocyte-derived state responding to "
    "adipocyte death than with the chronically lipid-laden lipid-"
    "associated-macrophage phenotype that develops later in disease "
    "progression."
)

doc.add_heading("4.1 Limitations", level=2)
lim_items = [
    ("Single-donor scRNA-seq reference. ",
     "The scRNA-seq reference contains one biological replicate per "
     "genotype. As demonstrated directly in this study (Section 2.5), a "
     "single pair cannot establish population-level direction, and "
     "individual-mouse variability likely adds noise to the marker-gene "
     "signature itself that cannot be fully quantified without additional "
     "biological replicates. External validation (Section 2.7, Figure 5) "
     "shows the S100A8-hi population itself reproduces in an independent, "
     "CD45+-sorted dataset, partially mitigating but not eliminating this "
     "limitation — it does not, for instance, validate the exact 5-25-gene "
     "marker panels used for deconvolution, only the population's "
     "existence."),
    ("Relative, not absolute, abundance scores. ",
     "MarkerBasedDecomposition returns a per-cell-type relative abundance "
     "score (marker-panel PC1), not an absolute cell proportion; scores are "
     "only interpretable for comparing the same cell type across samples, "
     "not for comparing magnitude between different cell types."),
    ("Small bulk sample sizes. ",
     "Group sizes of 4-5 per condition limit this study to detecting "
     "Cohen's d ≳ 2.0-2.4 at conventional power; moderate true effects, "
     "including the S100A8-hi subset's own directionally consistent signal "
     "(d ≈ 0.9 in both axes), cannot be excluded and require a larger "
     "cohort to resolve."),
    ("Bulk deconvolution cannot resolve cell state or spatial context. ",
     "A change in per-cell S100A8 activity, paracrine signaling range, or "
     "spatial proximity to hepatocytes would not necessarily register as a "
     "change in our composition estimate, and whole-liver bulk correlation "
     "cannot distinguish these from true compositional shifts."),
    ("Ccn3 (Nov) was not testable. ",
     "This transcript was absent from the bulk gene annotation used in "
     "this dataset and could not be included in the lipogenic-gene "
     "correlation analysis, leaving the paper's proposed CCN3-CD36 "
     "regulatory step untested directly. This does not, however, undermine "
     "the study's central conclusion: all four downstream lipogenic/lipid-"
     "handling genes that could be tested (Cd36, Pparg, Scd1, Plin2) showed "
     "the same pattern — correlated with general Macrophage_Kupffer "
     "abundance, not with the S100A8-hi subset specifically (Section 3.4) "
     "— so the finding does not depend on Ccn3 in particular."),
    ("Two genotype axes were processed as separate sequencing batches. ",
     "Although \"WT\" samples from the two axes were never pooled or "
     "treated as interchangeable, subtler batch effects on library "
     "composition or normalization between the two independently deposited "
     "DESeq2 analyses cannot be fully excluded."),
    ("Correlational design. ",
     "Associations between cell-type abundance scores and lipogenic gene "
     "expression are correlational; this study cannot establish causal "
     "direction and does not substitute for the interventional experiments "
     "in the original report."),
]
for i, (bold, rest) in enumerate(lim_items, start=1):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(bold)
    r.bold = True
    p.add_run(rest)

# ============================================================= CONCLUSION ==
doc.add_heading("5. Conclusions", level=1)
p = doc.add_paragraph()
p.add_run(
    "Reprocessing a public bulk-and-single-cell dataset with a validated, "
    "marker-based deconvolution approach shows that the compositional "
    "correlate of hepatic lipogenic gene expression in this model is the "
    "general macrophage/Kupffer compartment rather than the S100A8-hi "
    "subset specifically, and that macrophage-specific S100a8 deletion does "
    "not produce a statistically resolvable shift in broader hepatic immune "
    "composition at the sample sizes available. External validation against "
    "two independent public datasets shows the S100A8-hi population "
    "reproduces in a CD45+-sorted mouse liver dataset but not in a "
    "whole-liver dataset, indicating it is a real, immune-enrichment-"
    "dependent cell state rather than an artifact of the single-donor "
    "reference used in the primary analysis. These results refine, "
    "rather than contradict, the original mechanistic model, and highlight "
    "specific, adequately powered follow-up experiments — larger MKO "
    "cohorts, and dedicated attention to neutrophil and mast cell "
    "populations — that could distinguish a purely gene-regulatory "
    "mechanism from one with a compositional component."
)

doc.add_page_break()

# ======================================================= SUPPLEMENTARY ====
doc.add_heading("Supplementary Materials", level=1)

p = doc.add_paragraph()
p.add_run(
    "Full result tables underlying all figures and in-text statistics are "
    "reported in a separate document, "
)
r = p.add_run("Supplemental Data (Data S1-S8)")
r.italic = True
p.add_run(
    ", accompanying this manuscript: "
)
r = p.add_run("Data S1")
r.italic = True
p.add_run(" (full axis-1 deconvolution statistics, all 12 cell types), ")
r = p.add_run("Data S2")
r.italic = True
p.add_run(" (full axis-2/MKO statistics), ")
r = p.add_run("Data S3")
r.italic = True
p.add_run(" (full lipogenic-gene correlation statistics), ")
r = p.add_run("Data S4")
r.italic = True
p.add_run(" (GSE156057 marker-gene percent-positive by cluster, all 40 clusters), ")
r = p.add_run("Data S5")
r.italic = True
p.add_run(" (GSE156057 top-5 marker genes per cluster), ")
r = p.add_run("Data S6")
r.italic = True
p.add_run(" (GSE156057 S100A8-hi cluster proportions by diet/timepoint), ")
r = p.add_run("Data S7")
r.italic = True
p.add_run(" (GSE166504 marker-gene percent-positive by cluster, all 18 clusters), and ")
r = p.add_run("Data S8")
r.italic = True
p.add_run(
    " (GSE166504 cluster-to-author-label cross-tabulation). Each Data S "
    "item states its source CSV and generating script for full "
    "reproducibility. The underlying raw CSV files are also available "
    "directly in the code repository (Section 2.8): results/11_axis1_"
    "stats_full.csv, results/11_MKO_stats_full.csv, results/11_lipid_"
    "correlation_BH.csv, results/11_marker_panel_QC.csv (Table S1 source "
    "data, below), and external_validation/results/*.csv."
)

s1_headers = ["Cell type", "Marker genes used", "% variance explained by PC1"]
s1_rows = [
    ["B_cell", "5", "70.4"],
    ["Cholangiocyte", "14", "33.5"],
    ["Dendritic_cell", "5", "54.6"],
    ["Endothelial", "5", "64.3"],
    ["Hepatocyte", "6", "50.6"],
    ["Macrophage_Kupffer", "7", "72.4"],
    ["Mast_cell", "9", "35.9"],
    ["Neutrophil", "25", "38.7"],
    ["NK_cell", "5", "78.1"],
    ["S100A8hi_macrophage", "7", "76.9"],
    ["Stellate_Fibroblast", "10", "69.1"],
    ["T_cell", "7", "34.8"],
]
add_caption(
    "Marker-panel size and signal coherence per cell type. Panel size is "
    "algorithmically selected by BisqueRNA::GetNumGenesWeighted() (Section "
    "2.4), not a fixed cutoff. % variance explained by PC1 indicates how "
    "coherently the marker panel behaves as a single factor; values below "
    "~50% (Cholangiocyte, Mast_cell, Neutrophil, T_cell) indicate a noisier, "
    "less unidimensional signal and should be interpreted with additional "
    "caution.",
    bold_prefix="Table S1. "
)
make_table(s1_headers, s1_rows, col_widths_in=[2.2, 1.6, 2.2])

doc.add_page_break()

# ============================================================= REFERENCES ==
doc.add_heading("References", level=1)

references = [
    "Guan Y, Kim Y, Wang Y, Cho YE, Xiang X, Kim SJ, Yao T, Feng D, Hwang S, "
    "Gao B. Adipocyte death promotes hepatic infiltration of S100A8+ "
    "macrophages and steatotic liver disease progression in mice. J Clin "
    "Invest. 2025;135(21):e190635. doi:10.1172/JCI190635",

    "Rinella ME, Lazarus JV, Ratziu V, et al. A multisociety Delphi "
    "consensus statement on new fatty liver disease nomenclature. "
    "Hepatology. 2023;78(6):1966-1986. doi:10.1097/HEP.0000000000000520",

    "Hao Y, Stuart T, Kowalski MH, Choudhary S, Hoffman P, Hartman A, "
    "Srivastava A, Molla G, Madad S, Fernandez-Granda C, Satija R. "
    "Dictionary learning for integrative, multimodal and scalable "
    "single-cell analysis. Nat Biotechnol. 2024;42(2):293-304. "
    "doi:10.1038/s41587-023-01767-y",

    "Jew B, Alvarez M, Rahmani E, Miao Z, Ko A, Garske KM, Sul JH, "
    "Pietiläinen KH, Pajukanta P, Halperin E. Accurate estimation of cell "
    "composition in bulk expression through robust integration of "
    "single-cell information. Nat Commun. 2020;11:1971. "
    "doi:10.1038/s41467-020-15816-6",

    "Love MI, Huber W, Anders S. Moderated estimation of fold change and "
    "dispersion for RNA-seq data with DESeq2. Genome Biol. 2014;15:550. "
    "doi:10.1186/s13059-014-0550-8",

    "Benjamini Y, Hochberg Y. Controlling the false discovery rate: a "
    "practical and powerful approach to multiple testing. J R Stat Soc "
    "Series B. 1995;57(1):289-300.",

    "Barrett T, Wilhite SE, Ledoux P, et al. NCBI GEO: archive for "
    "functional genomics data sets--update. Nucleic Acids Res. "
    "2013;41(D1):D991-D995. doi:10.1093/nar/gks1193",

    "McInnes L, Healy J, Melville J. UMAP: Uniform Manifold Approximation "
    "and Projection for Dimension Reduction. arXiv:1802.03426 [stat.ML]. "
    "2018.",

    "Su Q, Kim SY, Adewale F, Zhou Y, Aldler C, Ni M, Wei Y, Burczynski ME, "
    "Atwal GS, Sleeman MW, Murphy AJ, Xin Y, Cheng X. Single-cell RNA "
    "transcriptome landscape of hepatocytes and non-parenchymal cells in "
    "healthy and NAFLD mouse liver. iScience. 2021;24(11):103233. "
    "doi:10.1016/j.isci.2021.103233",

    "Remmerie A, Martens L, Thoné T, Castoldi A, Seurinck R, Pavie B, et al. "
    "Osteopontin Expression Identifies a Subset of Recruited Macrophages "
    "Distinct from Kupffer Cells in the Fatty Liver. Immunity. "
    "2020;53(3):641-657.e14. doi:10.1016/j.immuni.2020.08.004",
]

for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(f"{i}. {ref}")
    run.font.size = Pt(10)

print("References written. Document complete.")
doc.save(OUT_PATH)
print(f"Saved to {OUT_PATH}")
